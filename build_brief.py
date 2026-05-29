#!/usr/bin/env python3
"""
Captello Onboarding Prep Brief builder (python-docx). Internal, confidential.
Usage:
    pip3 install python-docx
    python3 build_brief.py content.json out.docx
content.json schema:
{
  "client": "Philips",
  "call_type": "returning",
  "owner": "Amanda",
  "status_line": "Returning / already-kicked-off client · 2 working sessions remaining · 5-slide deck.",
  "snapshot": "One paragraph account snapshot.",
  "package": ["ULC x15", "QuickScan", "Salesforce integration"],   # purchased products/services
  "contacts": [["Caroline Salter","Primary contact — addressee on the pre-call email"], ...],
  "watchouts": ["Confirm Nadine's email ...", "..."],
  "tickets": [["#38416","Access issue","Waiting on customer"], ...],   # [] to omit section
  "email": {"to":"Caroline Salter","cc":"Jennifer Bourassa, Tanja Schaefer"},
  "agenda": ["Recap progress ...", "..."]
}
Logo logos/captello-horizontal-blacktext-transparent.png must exist relative to CWD.
"""
import json, sys, os
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

RED="FF0000"; BLACK="000000"; GRAY="777777"; LGRID="DDDDDD"; FONT="Roboto"
LOGO="logos/captello-horizontal-blacktext-transparent.png"
FOOTER=("\u00A9 COPYRIGHT | LEAD LIAISON | CAPTELLO | ALL RIGHTS RESERVED | "
        "13101 PRESTON ROAD STE 110 - 159 DALLAS, TX 75240 | 888.399.6430 | +44 20 38074910")

def shade(cell,fill):
    tcPr=cell._tc.get_or_add_tcPr(); sh=OxmlElement('w:shd')
    sh.set(qn('w:val'),'clear'); sh.set(qn('w:fill'),fill); tcPr.append(sh)
def bottom_border(p,color,sz):
    pPr=p._p.get_or_add_pPr(); pbdr=OxmlElement('w:pBdr'); b=OxmlElement('w:bottom')
    b.set(qn('w:val'),'single'); b.set(qn('w:sz'),str(sz)); b.set(qn('w:space'),'3'); b.set(qn('w:color'),color)
    pbdr.append(b); pPr.append(pbdr)
def run(p,text,size,color,bold=False):
    r=p.add_run(text); r.font.name=FONT; r.font.size=Pt(size); r.font.bold=bold; r.font.color.rgb=RGBColor.from_string(color)
    rpr=r._element.get_or_add_rPr(); rf=rpr.find(qn('w:rFonts'))
    if rf is None: rf=OxmlElement('w:rFonts'); rpr.append(rf)
    rf.set(qn('w:ascii'),FONT); rf.set(qn('w:hAnsi'),FONT)
    return r
def h2(doc,text):
    p=doc.add_paragraph(); p.paragraph_format.space_before=Pt(14); p.paragraph_format.space_after=Pt(7)
    run(p,text,13,BLACK,True); bottom_border(p,RED,14); return p
def body(doc,text,after=6):
    p=doc.add_paragraph(); p.paragraph_format.space_after=Pt(after); run(p,text,10.5,"262626"); return p

def table(doc,headers,rows,widths,status_col=None):
    t=doc.add_table(rows=1,cols=len(headers)); t.autofit=False
    t.style='Table Grid'
    hdr=t.rows[0].cells
    for i,h in enumerate(headers):
        hdr[i].width=Inches(widths[i]); shade(hdr[i],BLACK)
        p=hdr[i].paragraphs[0]; run(p,h,9.5,"FFFFFF",True)
    for row in rows:
        cells=t.add_row().cells
        for i,val in enumerate(row):
            cells[i].width=Inches(widths[i]); cells[i].vertical_alignment=WD_ALIGN_VERTICAL.CENTER
            p=cells[i].paragraphs[0]
            is_status = status_col is not None and i==status_col
            if is_status: shade(cells[i],"FCE9E9")
            run(p,val,9.5,"B30000" if is_status else "262626", bold=(i==0 or is_status))
    return t

def build(c,outpath):
    doc=Document()
    st=doc.styles['Normal']; st.font.name=FONT; st.font.size=Pt(10.5)
    sec=doc.sections[0]
    sec.top_margin=Inches(0.82); sec.bottom_margin=Inches(0.82); sec.left_margin=Inches(1); sec.right_margin=Inches(1)

    # header with logo + rule
    hp=sec.header.paragraphs[0]
    if os.path.exists(LOGO): hp.add_run().add_picture(LOGO,width=Inches(1.6))
    bottom_border(hp,LGRID,6)
    # footer
    fp=sec.footer.paragraphs[0]; fp.alignment=WD_ALIGN_PARAGRAPH.CENTER
    pPr=fp._p.get_or_add_pPr(); pbdr=OxmlElement('w:pBdr'); tb=OxmlElement('w:top')
    tb.set(qn('w:val'),'single'); tb.set(qn('w:sz'),'6'); tb.set(qn('w:space'),'4'); tb.set(qn('w:color'),LGRID); pbdr.append(tb); pPr.append(pbdr)
    run(fp,FOOTER,6.5,GRAY)

    # eyebrow + title
    p=doc.add_paragraph(); p.paragraph_format.space_after=Pt(2); run(p,"CONFIDENTIAL  \u00b7  INTERNAL USE ONLY",8,RED,True)
    p=doc.add_paragraph(); p.paragraph_format.space_after=Pt(3); run(p,f"{c['client']} — Onboarding Prep Brief",20,BLACK,True); bottom_border(p,RED,14)
    p=doc.add_paragraph(); p.paragraph_format.space_before=Pt(6); p.paragraph_format.space_after=Pt(3)
    run(p,"Status:  ",10.5,BLACK,True); run(p,c.get("status_line",""),10.5,"262626")

    h2(doc,"Account snapshot")
    body(doc,c.get("snapshot",""))
    if c.get("contacts"):
        table(doc,["Contact","Role in this engagement"],c["contacts"],[3.1,5.4])

    if c.get("package"):
        h2(doc,"Purchased products & services")
        body(doc,"Per the contract on file:",after=5)
        for sku in c["package"]:
            p=doc.add_paragraph(style='List Bullet'); p.paragraph_format.space_after=Pt(4); run(p,sku,10.5,"262626")

    h2(doc,f"Watch-outs for {c.get('owner','the onboarding specialist')}")
    for w in c.get("watchouts",[]):
        p=doc.add_paragraph(style='List Bullet'); p.paragraph_format.space_after=Pt(4); run(p,w,10.5,"262626")

    if c.get("tickets"):
        h2(doc,"Open Pylon tickets")
        body(doc,"Tickets below remain open. Raise the oldest on the next call.",after=7)
        table(doc,["Ticket","Issue","Status"],c["tickets"],[1.6,4.2,2.7],status_col=2)

    h2(doc,"Pre-call email draft")
    body(doc,f"A pre-call email is drafted and saved to Gmail. Draft only — review and send as {c.get('owner','the rep')}; nothing is sent automatically.")
    em=c.get("email",{})
    for label,val in [("To",em.get("to","")),("Cc",em.get("cc",""))]:
        if val:
            p=doc.add_paragraph(style='List Bullet'); p.paragraph_format.space_after=Pt(3); run(p,f"{label}: {val}",10.5,"262626")

    h2(doc,"Recommended agenda for the next call")
    for i,a in enumerate(c.get("agenda",[]),1):
        p=doc.add_paragraph(style='List Number'); p.paragraph_format.space_after=Pt(3); run(p,a,10.5,"262626")

    doc.save(outpath); print("WROTE",outpath)

if __name__=="__main__":
    c=json.load(open(sys.argv[1])) if len(sys.argv)>1 else {"client":"Sample","call_type":"returning"}
    out=sys.argv[2] if len(sys.argv)>2 else "brief.docx"
    build(c,out)
